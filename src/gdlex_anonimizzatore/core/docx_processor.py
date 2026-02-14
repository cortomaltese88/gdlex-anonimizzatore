from __future__ import annotations

from dataclasses import replace
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from gdlex_anonimizzatore.core.models import EntityFinding, EntityType, Settings
from gdlex_anonimizzatore.core.recognizers import detect_entities


def _iter_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def extract_docx_text_and_findings(path: Path, settings: Settings) -> tuple[str, list[EntityFinding]]:
    document = Document(str(path))
    text_parts: list[str] = []
    findings: list[EntityFinding] = []
    cursor = 0

    for paragraph in _iter_paragraphs(document):
        text = paragraph.text
        text_parts.append(text)
        para_findings = detect_entities(text, settings)
        for finding in para_findings:
            findings.append(replace(finding, start=finding.start + cursor, end=finding.end + cursor))
        cursor += len(text) + 1

    return "\n".join(text_parts), findings


def anonymize_docx_file(
    input_path: Path,
    output_path: Path,
    settings: Settings,
    findings: list[EntityFinding],
) -> tuple[int, int]:
    """Run-level safe replacement: replace only when a whole match is inside one run."""
    document = Document(str(input_path))

    overrides: dict[tuple[EntityType, str], str] = {}
    disabled: set[tuple[EntityType, str]] = set()
    for finding in findings:
        key = (finding.entity_type, finding.value)
        if not finding.anonymize:
            disabled.add(key)
            continue
        overrides[key] = finding.replacement

    applied = 0
    skipped = 0

    for paragraph in _iter_paragraphs(document):
        para_text = paragraph.text
        if not para_text.strip():
            continue
        para_findings = detect_entities(para_text, settings)
        para_findings.sort(key=lambda item: len(item.value), reverse=True)

        for finding in para_findings:
            key = (finding.entity_type, finding.value)
            if key in disabled:
                continue
            replacement = overrides.get(key, finding.replacement)

            replaced_here = False
            for run in paragraph.runs:
                if finding.value in run.text:
                    run.text = run.text.replace(finding.value, replacement)
                    replaced_here = True
                    applied += 1
                    break

            if not replaced_here:
                skipped += 1

    document.save(str(output_path))
    return applied, skipped
