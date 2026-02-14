import pytest

pytest.importorskip("docx")

from pathlib import Path

from docx import Document

from gdlex_anonimizzatore.core.models import EntityType, Settings
from gdlex_anonimizzatore.core.docx_processor import anonymize_docx_file, extract_docx_text_and_findings


def test_docx_basic_replacement(tmp_path: Path) -> None:
    in_path = tmp_path / "input.docx"
    out_path = tmp_path / "input_anonimizzato.docx"

    doc = Document()
    doc.add_paragraph("ROSSI Mario scrive a mario.rossi@example.com")
    doc.add_paragraph("La società Eni S.p.A.")
    doc.save(in_path)

    settings = Settings()
    _, findings = extract_docx_text_and_findings(in_path, settings)
    anonymize_docx_file(in_path, out_path, settings, findings)

    out_doc = Document(out_path)
    out_text = "\n".join(p.text for p in out_doc.paragraphs)
    assert "Tizio" in out_text
    assert "[EMAIL]" in out_text
    assert "Alfa" in out_text


def test_docx_table_replacement(tmp_path: Path) -> None:
    in_path = tmp_path / "table.docx"
    out_path = tmp_path / "table_anonimizzato.docx"

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Eni S.p.A."
    doc.save(in_path)

    settings = Settings()
    _, findings = extract_docx_text_and_findings(in_path, settings)
    anonymize_docx_file(in_path, out_path, settings, findings)

    out_doc = Document(out_path)
    assert "Alfa" in out_doc.tables[0].cell(0, 0).text


def test_docx_cross_run_safety_skip(tmp_path: Path) -> None:
    in_path = tmp_path / "split.docx"
    out_path = tmp_path / "split_anonimizzato.docx"

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("ROSSI ")
    p.add_run("Mario")
    doc.save(in_path)

    settings = Settings()
    _, findings = extract_docx_text_and_findings(in_path, settings)
    applied, skipped = anonymize_docx_file(in_path, out_path, settings, findings)

    out_doc = Document(out_path)
    assert out_doc.paragraphs[0].text == "ROSSI Mario"
    assert skipped >= 1
    assert applied == 0
    assert out_path.exists()


def test_docx_detects_all_supported_entities(tmp_path: Path) -> None:
    in_path = tmp_path / "entities.docx"
    doc = Document()
    doc.add_paragraph("Mario ROSSI CF RSSMRA85M01H501U PIVA 12345678901 Eni S.p.A. foo@example.com")
    doc.save(in_path)

    settings = Settings()
    _, findings = extract_docx_text_and_findings(in_path, settings)
    kinds = {f.entity_type for f in findings}
    assert EntityType.PERSONA in kinds
    assert EntityType.CODICE_FISCALE in kinds
    assert EntityType.PARTITA_IVA in kinds
    assert EntityType.SOCIETA in kinds
    assert EntityType.EMAIL in kinds
